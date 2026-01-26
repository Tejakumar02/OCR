import streamlit as st
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import time
import base64
import io
import cv2
import numpy as np
from tabula.io import read_pdf  # Existing tabula import
import pdfplumber  # New library for enhanced table extraction
from concurrent.futures import ThreadPoolExecutor
import zipfile  # For packaging extracted content into a zip file
from docx import Document  # For creating Word documents
from docx.shared import Inches  # For image sizing in Word documents
import re  # Import the re module for regular expressions
import pandas as pd  # For handling CSV files

# Streamlit App
st.title("PDF to Image, Text, Table, and Figure Extractor (CPU CODE) 📑")

# File uploader updated to allow multiple file types
uploaded_file = st.file_uploader("Upload a PDF, Image, Text/MD, or DOCX file 📂", type=["pdf", "png", "jpg", "jpeg", "txt", "md", "docx"])

# Function to extract text from an image
def extract_text(img):
    return pytesseract.image_to_string(img, config='--oem 1 --psm 3 -l eng') #oem 1 for more accurate
                                                                            #oem 3 for faster result

# Function to extract tables from a PDF page
def extract_tables(pdf_path, page_number):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            page = pdf.pages[page_number - 1]
            tables = page.extract_tables()
            return tables
    except Exception as e:
        return f"No tables found on Page {page_number} or an error occurred: {e}"

# Function to detect figures in an image
def detect_figures(img_np):
    gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    figures = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        if w > 100 and h > 100:  # Filter small objects
            figure_crop = img_np[y:y+h, x:x+w]
            if figure_crop.size > 0:  # Ensure valid figure
                figures.append(figure_crop)
    return figures

# Function to sanitize text for XML compatibility
def sanitize_text(text):
    # Remove NULL bytes and control characters
    sanitized_text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    return sanitized_text

# Function to convert formulas to LaTeX format
def convert_formulas_to_latex(text):
    # Example: Replace simple patterns with LaTeX
    text = re.sub(r'(\d+)\^(\d+)', r'\1^{\2}', text)  # Exponents
    text = re.sub(r'(\d+)_(\d+)', r'\1_{\2}', text)  # Subscripts
    text = re.sub(r'sqrt\(([^)]+)\)', r'\\sqrt{\1}', text)  # Square roots
    text = re.sub(r'pi', r'\\pi', text)  # Greek letters
    text = re.sub(r'alpha', r'\\alpha', text)  # Greek letters
    text = re.sub(r'beta', r'\\beta', text)  # Greek letters
    return text

# URL extraction pattern
URL_PATTERN = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'

# Apply custom CSS for UI enhancement
st.markdown("""
<style>
    .stButton > button {
        background-color: #4CAF50;
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 5px;
        cursor: pointer;
        font-size: 16px;
        margin: 5px;
        transition: background-color 0.3s;
    }
    .stButton > button:hover {
        background-color: #45a049;
    }
    .stFileUploader > div > div {
        border: 2px dashed #4CAF50;
        border-radius: 10px;
        padding: 20px;
        background-color: #f9f9f9;
    }
    .stProgress > div > div {
        background-color: #4CAF50;
    }
    h1, h2, h3 {
        color: #333;
    }
    .sidebar .sidebar-content {
        background-color: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
    }
</style>
""", unsafe_allow_html=True)

if uploaded_file:
    file_ext = uploaded_file.name.split('.')[-1].lower()
    if file_ext == "pdf":
        # Start timing
        start_time = time.time()

        # Convert uploaded file into a BytesIO object
        file_stream = io.BytesIO(uploaded_file.read())  # Convert to BytesIO

        # Open PDF with fitz using file_stream.read() to avoid TypeError
        pdf_document = fitz.open("pdf", file_stream.read())  # Use the PDF bytes directly
        total_pages = len(pdf_document)
        images = []

        # Progress bar for image extraction
        st.write("Extracting images from PDF... 🖼️")
        progress_bar = st.progress(0)
        for page_number in range(total_pages):
            page = pdf_document[page_number]
            pix = page.get_pixmap()  # type: ignore  # Suppress Pyright error
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

            # Convert image to base64 for display
            buffer = io.BytesIO()
            img.save(buffer, format="PNG")
            buffer.seek(0)
            img_base64 = base64.b64encode(buffer.read()).decode('utf-8')

            # Store the base64 image for later use
            images.append(img_base64)

            # Update progress bar
            progress_bar.progress((page_number + 1) / total_pages)

        pdf_document.close()

        # Text, Table, and Figure extraction for all pages
        st.write("### Extract Text, Tables, and Figures from All Pages 📚")

        if st.button("Extract All Content 📥"):
            extracted_content = {}
            text_content = ""
            table_content = []
            figure_content = []
            links_content = []
            text_start_time = time.time()
            text_extraction_time = 0  # Initialize text extraction time

            # Extract PDF links
            uploaded_file.seek(0)
            with fitz.open("pdf", uploaded_file.read()) as pdf_document_links:
                links_per_page = []
                for page in pdf_document_links:
                    links = page.get_links()  # type: ignore
                    uri_list = [link.get('uri', '') for link in links if 'uri' in link]
                    links_per_page.append(uri_list)

            # Progress bar for content extraction
            st.write("Extracting text, tables, and figures... 🔍")
            progress_bar = st.progress(0)
            for page_number, image_data in enumerate(images, start=1):
                img = Image.open(io.BytesIO(base64.b64decode(image_data)))
                img_np = np.array(img)

                # Use ThreadPoolExecutor to process text, tables, and figures in parallel
                with ThreadPoolExecutor() as executor:
                    # Submit tasks
                    future_text = executor.submit(extract_text, img)
                    future_tables = executor.submit(extract_tables, uploaded_file, page_number)
                    future_figures = executor.submit(detect_figures, img_np)

                    # Wait for all tasks to complete
                    extracted_text = future_text.result()
                    tables = future_tables.result()
                    detected_figures = future_figures.result()

                # Convert formulas in the extracted text to LaTeX format
                extracted_text = convert_formulas_to_latex(extracted_text)

                # Extract links from OCR text
                text_urls = re.findall(URL_PATTERN, extracted_text)
                structured_links = links_per_page[page_number - 1]
                all_links = list(set(structured_links + text_urls))

                # Calculate text extraction time for this page
                text_extraction_time += time.time() - text_start_time

                # Store content per page
                extracted_content[page_number] = {
                    "text": extracted_text,
                    "tables": tables,
                    "figures": detected_figures,
                    "links": all_links
                }

                # Append extracted content to overall outputs
                text_content += f"### Page {page_number}\n{extracted_text}\n\n"
                table_content.append((page_number, tables))
                figure_content.append((page_number, detected_figures))
                links_content.append((page_number, all_links))

                # Update progress bar
                progress_bar.progress(page_number / total_pages)

            text_processing_time = time.time() - text_start_time

            # Preview Extracted Content
            st.write("### Extracted Text 📝")
            st.markdown(text_content)  # Use st.markdown to render LaTeX

            st.write("### Extracted Tables 📊")
            for page_number, tables in table_content:
                if isinstance(tables, list):
                    for idx, table in enumerate(tables, start=1):
                        st.write(f"Table {idx} on Page {page_number}:")
                        st.dataframe(table)
                else:
                    st.write(tables)

            st.write("### Extracted Figures 🖼️")
            for page_number, detected_figures in figure_content:
                for idx, figure_crop in enumerate(detected_figures, start=1):
                    fig_img = Image.fromarray(figure_crop)
                    st.image(fig_img, caption=f"Figure {idx} from Page {page_number}", use_container_width=True)

            st.write("### Extracted Links 🔗")
            for page_number, links in links_content:
                if links:
                    st.write(f"Page {page_number}:")
                    for link in links:
                        st.write(f"- {link}")

            # Display extraction times
            st.write(f"Text Extraction Time: {text_extraction_time:.2f} seconds ⏱️")
            st.write(f"Total Content Extraction Time: {text_processing_time:.2f} seconds ⏱️")

            # Download Buttons
            st.download_button(
                "Download Extracted Text 📄",
                text_content,
                file_name="extracted_text.md",
                mime="text/markdown",
            )

            # Save tables as CSV files in a zip
            table_zip = io.BytesIO()
            with zipfile.ZipFile(table_zip, "w") as zipf:
                for page_number, tables in table_content:
                    if isinstance(tables, list):
                        for idx, table in enumerate(tables, start=1):
                            # Convert the table to a pandas DataFrame
                            df = pd.DataFrame(table)
                            # Save the DataFrame as a CSV file in the zip
                            csv_buffer = io.StringIO()
                            df.to_csv(csv_buffer, index=False)
                            csv_buffer.seek(0)
                            zipf.writestr(f"Page_{page_number}Table{idx}.csv", csv_buffer.getvalue())
                    else:
                        zipf.writestr(f"Page_{page_number}_tables_error.txt", tables)
            table_zip.seek(0)
            st.download_button(
                "Download Extracted Tables 📊",
                table_zip,
                file_name="extracted_tables.zip",
                mime="application/zip",
            )

            # Save figures as PNG files in a zip
            figure_zip = io.BytesIO()
            with zipfile.ZipFile(figure_zip, "w") as zipf:
                for page_number, detected_figures in figure_content:
                    for idx, figure_crop in enumerate(detected_figures, start=1):
                        fig_img = Image.fromarray(figure_crop)
                        buffer = io.BytesIO()
                        fig_img.save(buffer, format="PNG")
                        buffer.seek(0)
                        zipf.writestr(f"Page_{page_number}Figure{idx}.png", buffer.read())
            figure_zip.seek(0)
            st.download_button(
                "Download Extracted Figures 🖼️",
                figure_zip,
                file_name="extracted_figures.zip",
                mime="application/zip",
            )

            # Save links as text files in a zip
            links_zip = io.BytesIO()
            with zipfile.ZipFile(links_zip, "w") as zipf:
                for page_number, links in links_content:
                    if links:
                        content = "\n".join(links)
                        zipf.writestr(f"Page_{page_number}_links.txt", content)
            links_zip.seek(0)
            st.download_button(
                "Download Extracted Links 🔗",
                links_zip,
                file_name="extracted_links.zip",
                mime="application/zip",
            )

            # "Download All" Functionality in DOCX format
            all_content_doc = Document()

            for page_number, content in extracted_content.items():
                # Add page number as a heading
                all_content_doc.add_heading(f"Page {page_number}", level=1)
                
                # Add text content (sanitized)
                if content["text"]:
                    sanitized_text = sanitize_text(content["text"])
                    all_content_doc.add_paragraph(sanitized_text)
                
                # Add tables (sanitized)
                if isinstance(content["tables"], list):
                    for idx, table in enumerate(content["tables"], start=1):
                        all_content_doc.add_heading(f"Table {idx}", level=2)
                        # Convert the table to a pandas DataFrame
                        df = pd.DataFrame(table)
                        # Convert the DataFrame to a formatted table in the Word document
                        t = all_content_doc.add_table(df.shape[0] + 1, df.shape[1])
                        # Add the header row
                        for j in range(df.shape[1]):
                            header_value = str(df.columns[j]) if j < len(df.columns) else ""
                            t.cell(0, j).text = header_value
                        # Add the rest of the data
                        for i in range(df.shape[0]):
                            for j in range(df.shape[1]):
                                cell_value = str(df.values[i, j]) if df.values[i, j] is not None else ""
                                t.cell(i + 1, j).text = cell_value
                else:
                    # Sanitize the error message or non-table content
                    sanitized_tables = sanitize_text(content["tables"])
                    all_content_doc.add_paragraph(sanitized_tables)
                
                # Add figures
                for idx, figure_crop in enumerate(content["figures"], start=1):
                    all_content_doc.add_heading(f"Figure {idx}", level=2)
                    fig_img = Image.fromarray(figure_crop)
                    buffer = io.BytesIO()
                    fig_img.save(buffer, format="PNG")
                    buffer.seek(0)
                    all_content_doc.add_picture(buffer, width=Inches(4.0))  # Adjust width as needed
                
                # Add links
                if content["links"]:
                    all_content_doc.add_heading("Links", level=2)
                    for link in content["links"]:
                        all_content_doc.add_paragraph(link)
            
            # Save the document to a BytesIO object
            doc_buffer = io.BytesIO()
            all_content_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.download_button(
                "Download All Content 📦",
                doc_buffer,
                file_name="extracted_content.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    elif file_ext in ["png", "jpg", "jpeg"]:
        # Process image files
        img = Image.open(uploaded_file)
        st.image(img, caption="Uploaded Image 🖼️", use_container_width=True)
        extracted_text = extract_text(img)
        urls = re.findall(URL_PATTERN, extracted_text)
        st.write("### Extracted Text from Image 📝")
        st.write(extracted_text)
        st.write("### Extracted Links from Image 🔗")
        st.write(urls)
        if urls:
            links_text = "\n".join(urls)
            st.download_button(
                "Download Links 🔗",
                links_text,
                file_name="extracted_links.txt",
                mime="text/plain",
            )
    elif file_ext in ["txt", "md"]:
        # Process text and markdown files
        content = uploaded_file.read().decode("utf-8")
        urls = re.findall(URL_PATTERN, content)
        st.write("### File Content 📝")
        st.text_area("Content", content, height=300)
        st.write("### Extracted Links 🔗")
        st.write(urls)
        if urls:
            links_text = "\n".join(urls)
            st.download_button(
                "Download Links 🔗",
                links_text,
                file_name="extracted_links.txt",
                mime="text/plain",
            )
    elif file_ext == "docx":
        # Process docx files
        doc = Document(uploaded_file)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        content = "\n".join(fullText)
        urls = re.findall(URL_PATTERN, content)
        st.write("### DOCX Content 📝")
        st.text_area("Content", content, height=300)
        st.write("### Extracted Links 🔗")
        st.write(urls)
        if urls:
            links_text = "\n".join(urls)
            st.download_button(
                "Download Links 🔗",
                links_text,
                file_name="extracted_links.txt",
                mime="text/plain",
            )