import streamlit as st
from paddleocr import PaddleOCR  # New OCR engine
from PIL import Image
import fitz  # PyMuPDF
import time
import base64
import io
import cv2
import numpy as np
from tabula.io import read_pdf
import pdfplumber
from concurrent.futures import ThreadPoolExecutor
import zipfile
from docx import Document
from docx.shared import Inches
import re
import pandas as pd

# Custom UI styling and animation additions
st.set_page_config(page_title="Smart PDF Extractor", layout="wide")
st.markdown("""
    <style>
    .css-1d391kg {padding: 2rem;}
    .stButton>button {background-color: #4CAF50; color: white; font-weight: bold; border-radius: 10px; transition: all 0.3s ease;}
    .stButton>button:hover {background-color: #45a049; transform: scale(1.05);}
    .css-1cpxqw2, .css-1kyxreq {background-color: #f8f9fa; border-radius: 15px; padding: 10px; box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);}
    .block-container {padding-top: 2rem;}
    </style>
""", unsafe_allow_html=True)

# Initialize PaddleOCR (enable GPU if available)
ocr = PaddleOCR(use_angle_cls=True, lang="en")

st.title(" 📑 Smart PDF Extractor with OCR & AI")

uploaded_file = st.file_uploader(" 📤 Upload a PDF file 📂", type=["pdf"])

# Function to extract text using PaddleOCR
def extract_text(img):
    result = ocr.ocr(np.array(img), cls=True)
    extracted_text = "\n".join([line[1][0] for line in result[0]]) if result[0] else ""
    return extracted_text

# Function to extract tables
def extract_tables(pdf_path, page_number):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            page = pdf.pages[page_number - 1]
            tables = page.extract_tables()
            return tables
    except Exception as e:
        return f"No tables found on Page {page_number} or an error occurred: {e}"

# Function to detect figures
def detect_figures(img_np):
    gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    figures = [img_np[y:y+h, x:x+w] for x, y, w, h in [cv2.boundingRect(cnt) for cnt in contours] if w > 100 and h > 100]
    return figures

# Sanitize text for compatibility
def sanitize_text(text):
    return re.sub(r'[\x00-\x1F\x7F]', '', text)

# Convert formulas to LaTeX
def convert_formulas_to_latex(text):
    text = re.sub(r'(\d+)\^(\d+)', r'\1^{\2}', text)
    text = re.sub(r'(\d+)([0-9]+)', r'\1{\2}', text)
    text = re.sub(r'sqrt\(([^)]+)\)', r'\\sqrt{\1}', text)
    text = re.sub(r'pi', r'\\pi', text)
    return text

# URL extraction pattern
URL_PATTERN = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'

if uploaded_file:
    start_time = time.time()
    file_stream = io.BytesIO(uploaded_file.read())
    pdf_document = fitz.open("pdf", file_stream.read())
    total_pages = len(pdf_document)
    images = []

    st.subheader("🔍 Extracting images from PDF... 🖼️")
    progress_bar = st.progress(0)

    for page_number in range(total_pages):
        page = pdf_document[page_number]
        pix = page.get_pixmap() # type: ignore
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)
        img_base64 = base64.b64encode(buffer.read()).decode('utf-8')

        images.append(img_base64)
        progress_bar.progress((page_number + 1) / total_pages)

    pdf_document.close()

    st.markdown("---")
    st.header("🧠 Extract Text, Tables, and Figures from All Pages 📚")

    if st.button("🚀 Extract All Content 📥"):
        extracted_content = {}
        text_content = ""
        table_content = []
        figure_content = []
        links_content = []
        text_start_time = time.time()
        text_extraction_time = 0

        # Extract PDF links
        uploaded_file.seek(0)
        with fitz.open("pdf", uploaded_file.read()) as pdf_document_links:
            links_per_page = []
            for page in pdf_document_links:
                links = page.get_links() # type: ignore
                uri_list = [link.get('uri', '') for link in links if 'uri' in link]
                links_per_page.append(uri_list)

        st.info("Working through your PDF pages. Please wait... 🔍")
        progress_bar = st.progress(0)

        for page_number, image_data in enumerate(images, start=1):
            img = Image.open(io.BytesIO(base64.b64decode(image_data)))
            img_np = np.array(img)

            with ThreadPoolExecutor() as executor:
                future_text = executor.submit(extract_text, img)
                future_tables = executor.submit(extract_tables, uploaded_file, page_number)
                future_figures = executor.submit(detect_figures, img_np)

                extracted_text = future_text.result()
                tables = future_tables.result()
                detected_figures = future_figures.result()

            extracted_text = convert_formulas_to_latex(extracted_text)
            text_urls = re.findall(URL_PATTERN, extracted_text)
            structured_links = links_per_page[page_number - 1]
            all_links = list(set(structured_links + text_urls))

            text_extraction_time += time.time() - text_start_time

            extracted_content[page_number] = {
                "text": extracted_text,
                "tables": tables,
                "figures": detected_figures,
                "links": all_links
            }

            text_content += f"### Page {page_number}\n{extracted_text}\n\n"
            table_content.append((page_number, tables))
            figure_content.append((page_number, detected_figures))
            links_content.append((page_number, all_links))

            progress_bar.progress(page_number / total_pages)

        text_processing_time = time.time() - text_start_time

        st.subheader("📝 Extracted Text")
        st.markdown(text_content)

        st.subheader("📊 Extracted Tables")
        for page_number, tables in table_content:
            if isinstance(tables, list):
                for idx, table in enumerate(tables, start=1):
                    st.write(f"Table {idx} on Page {page_number}:")
                    st.dataframe(table)
            else:
                st.warning(tables)

        st.subheader("🖼️ Extracted Figures")
        for page_number, detected_figures in figure_content:
            for idx, figure_crop in enumerate(detected_figures, start=1):
                fig_img = Image.fromarray(figure_crop)
                st.image(fig_img, caption=f"Figure {idx} from Page {page_number}", use_container_width=True)

        st.subheader("🔗 Extracted Links")
        for page_number, links in links_content:
            if links:
                st.write(f"Page {page_number}:")
                for link in links:
                    st.write(f"- {link}")

        st.success(f"✅ Text Extraction Time: {text_extraction_time:.2f} seconds ⏱️")
        st.success(f"⏱ Total Content Extraction Time: {text_processing_time:.2f} seconds ⏱️")

        st.download_button(" 📥 Download Extracted Text 💾", text_content, "extracted_text.md", "text/markdown")

        table_zip = io.BytesIO()
        with zipfile.ZipFile(table_zip, "w") as zipf:
            for page_number, tables in table_content:
                if isinstance(tables, list):
                    for idx, table in enumerate(tables, start=1):
                        df = pd.DataFrame(table)
                        csv_buffer = io.StringIO()
                        df.to_csv(csv_buffer, index=False)
                        csv_buffer.seek(0)
                        zipf.writestr(f"Page_{page_number}Table{idx}.csv", csv_buffer.getvalue())
                else:
                    zipf.writestr(f"Page_{page_number}_tables_error.txt", tables)
        table_zip.seek(0)
        st.download_button("📥 Download Extracted Tables 📊", table_zip, "extracted_tables.zip", "application/zip")

        figure_zip = io.BytesIO()
        with zipfile.ZipFile(figure_zip, "w") as zipf:
            for page_number, detected_figures in figure_content:
                for idx, figure_crop in enumerate(detected_figures, start=1):
                    fig_img = Image.fromarray(figure_crop)
                    buffer = io.BytesIO()
                    fig_img.save(buffer, format="PNG")
                    buffer.seek(0)
                    zipf.writestr(f"Page_{page_number}Figure{idx}.png", buffer.read())
        figure_zip.seek(0)
        st.download_button("📥 Download Extracted Figures 🖼️", figure_zip, "extracted_figures.zip", "application/zip")

        links_zip = io.BytesIO()
        with zipfile.ZipFile(links_zip, "w") as zipf:
            for page_number, links in links_content:
                if links:
                    content = "\n".join(links)
                    zipf.writestr(f"Page_{page_number}_links.txt", content)
        links_zip.seek(0)
        st.download_button("📥 Download Extracted Links 🔗", links_zip, "extracted_links.zip", "application/zip")

        # "Download All" Functionality in DOCX format
        all_content_doc = Document()

        for page_number, content in extracted_content.items():
            # Add page number as a heading
            all_content_doc.add_heading(f"Page {page_number}", level=1)
            
            # Add text content (sanitized)
            if content["text"]:
                sanitized_text = sanitize_text(content["text"])
                all_content_doc.add_paragraph(sanitized_text)
            
            # Add tables (sanitized)
            if isinstance(content["tables"], list):
                for idx, table in enumerate(content["tables"], start=1):
                    all_content_doc.add_heading(f"Table {idx}", level=2)
                    # Convert the table to a pandas DataFrame
                    df = pd.DataFrame(table)
                    # Convert the DataFrame to a formatted table in the Word document
                    t = all_content_doc.add_table(df.shape[0] + 1, df.shape[1])
                    # Add the header row
                    for j in range(df.shape[1]):
                        header_value = str(df.columns[j]) if j < len(df.columns) else ""
                        t.cell(0, j).text = header_value
                    # Add the rest of the data
                    for i in range(df.shape[0]):
                        for j in range(df.shape[1]):
                            cell_value = str(df.values[i, j]) if df.values[i, j] is not None else ""
                            t.cell(i + 1, j).text = cell_value
            else:
                # Sanitize the error message or non-table content
                sanitized_tables = sanitize_text(content["tables"])
                all_content_doc.add_paragraph(sanitized_tables)
            
            # Add figures
            for idx, figure_crop in enumerate(content["figures"], start=1):
                all_content_doc.add_heading(f"Figure {idx}", level=2)
                fig_img = Image.fromarray(figure_crop)
                buffer = io.BytesIO()
                fig_img.save(buffer, format="PNG")
                buffer.seek(0)
                all_content_doc.add_picture(buffer, width=Inches(4.0))
            
            # Add links
            if content["links"]:
                all_content_doc.add_heading("Links", level=2)
                for link in content["links"]:
                    all_content_doc.add_paragraph(link)
        
        # Save the document to a BytesIO object
        doc_buffer = io.BytesIO()
        all_content_doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        st.download_button(
            "📥 Download All Content 📦",
            doc_buffer,
            file_name="extracted_content.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )